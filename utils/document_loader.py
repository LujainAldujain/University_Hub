"""Raw text extraction for the three supported university document formats.

Used by the Bronze Lakehouse loader to land the actual document content
(not just metadata) — and later reused by the RAG chunking stage.
"""
from __future__ import annotations

from pathlib import Path


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


_LOADERS = {
    "txt": load_txt,
    "docx": load_docx,
    "pdf": load_pdf,
}


def extract_text(file_path: str, file_type: str) -> str:
    """Extracts raw text from a document given its validated file_type."""
    loader = _LOADERS.get(file_type)
    if loader is None:
        raise ValueError(f"No text loader registered for file_type={file_type!r}")
    return loader(Path(file_path))
