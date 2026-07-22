"""One-off fixture generator: builds the DOCX and PDF sample university
documents used to exercise the ingestion pipeline (multi-format support) and,
later, the RAG pipeline. Run once with:

    .venv/Scripts/python.exe scripts/seed_sample_documents.py
"""
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

INCOMING = Path(__file__).resolve().parent.parent / "incoming_documents"
INCOMING.mkdir(exist_ok=True)


def build_cs_program_docx() -> None:
    doc = Document()
    doc.add_heading("Computer Science Program Requirements", level=1)
    doc.add_paragraph("Department of Computer Science — Academic Year 2026")

    doc.add_heading("1. Degree Requirements", level=2)
    doc.add_paragraph(
        "The Bachelor of Science in Computer Science requires 132 total "
        "credit hours, comprising University Core Curriculum, Major Core "
        "courses, Major Electives, and free electives."
    )

    doc.add_heading("2. Required Major Core Courses (42 credit hours)", level=2)
    for course in [
        "CS 101 — Introduction to Programming (3 credit hours)",
        "CS 102 — Data Structures and Algorithms (3 credit hours)",
        "CS 201 — Discrete Mathematics (3 credit hours)",
        "CS 210 — Computer Organization and Architecture (3 credit hours)",
        "CS 220 — Object-Oriented Software Design (3 credit hours)",
        "CS 301 — Database Systems (3 credit hours)",
        "CS 305 — Operating Systems (3 credit hours)",
        "CS 310 — Computer Networks (3 credit hours)",
        "CS 320 — Software Engineering (3 credit hours)",
        "CS 340 — Artificial Intelligence (3 credit hours)",
        "CS 401 — Senior Capstone Project I (3 credit hours)",
        "CS 402 — Senior Capstone Project II (3 credit hours)",
        "MATH 210 — Linear Algebra (3 credit hours)",
        "MATH 220 — Probability and Statistics (3 credit hours)",
    ]:
        doc.add_paragraph(course, style="List Bullet")

    doc.add_heading("3. Major Electives (12 credit hours)", level=2)
    doc.add_paragraph(
        "Students choose four courses from the approved elective list, which "
        "includes Machine Learning, Distributed Systems, Cybersecurity, Cloud "
        "Computing, Data Engineering, and Mobile Application Development."
    )

    doc.add_heading("4. Minimum GPA", level=2)
    doc.add_paragraph(
        "A minimum GPA of 2.0 in the major core is required to progress to "
        "the Senior Capstone sequence (CS 401 / CS 402)."
    )

    doc.add_heading("5. Internship Recommendation", level=2)
    doc.add_paragraph(
        "Students are strongly encouraged, though not required, to complete "
        "a summer internship between their third and fourth years."
    )

    out_path = INCOMING / "cs_program_requirements.docx"
    doc.save(str(out_path))
    print(f"wrote {out_path} ({out_path.stat().st_size} bytes)")


def build_attendance_policy_pdf() -> None:
    out_path = INCOMING / "attendance_policy.pdf"
    styles = getSampleStyleSheet()
    story = [
        Paragraph("University Attendance Policy", styles["Title"]),
        Spacer(1, 12),
        Paragraph("Office of the Registrar — Academic Year 2026", styles["Normal"]),
        Spacer(1, 12),
        Paragraph("1. General Attendance Requirement", styles["Heading2"]),
        Paragraph(
            "Students are expected to attend all scheduled class sessions. "
            "A student who misses more than 20% of the total scheduled "
            "sessions in a course, whether excused or unexcused, may be "
            "administratively withdrawn from that course with a grade of 'WF'.",
            styles["Normal"],
        ),
        Spacer(1, 8),
        Paragraph("2. Excused Absences", styles["Heading2"]),
        Paragraph(
            "Absences due to documented medical emergencies, bereavement, "
            "military service, or official university-sponsored events are "
            "excused, provided supporting documentation is submitted to the "
            "instructor within five business days of the absence.",
            styles["Normal"],
        ),
        Spacer(1, 8),
        Paragraph("3. Instructor Responsibilities", styles["Heading2"]),
        Paragraph(
            "Instructors must record attendance at every session and report "
            "any student who has not attended within the first two weeks of "
            "the semester to the Registrar for automatic administrative drop.",
            styles["Normal"],
        ),
        Spacer(1, 8),
        Paragraph("4. Appeals", styles["Heading2"]),
        Paragraph(
            "A student administratively withdrawn for excessive absence may "
            "appeal to the Dean of Students within 10 business days of "
            "notification, providing documentation for any disputed absences.",
            styles["Normal"],
        ),
    ]
    SimpleDocTemplate(str(out_path), pagesize=letter).build(story)
    print(f"wrote {out_path} ({out_path.stat().st_size} bytes)")


def build_bad_files() -> None:
    # Deliberately malformed inputs used to prove the DLQ / quarantine path.
    bad_ext = INCOMING / "corrupted_upload.exe"
    bad_ext.write_bytes(b"MZ\x90\x00\x03\x00\x00\x00 this is not a real document")
    print(f"wrote {bad_ext} ({bad_ext.stat().st_size} bytes) [disallowed extension]")

    empty_file = INCOMING / "empty_handbook.txt"
    empty_file.write_bytes(b"")
    print(f"wrote {empty_file} ({empty_file.stat().st_size} bytes) [zero-byte upload]")


if __name__ == "__main__":
    build_cs_program_docx()
    build_attendance_policy_pdf()
    build_bad_files()
